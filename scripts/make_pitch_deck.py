# -*- coding: utf-8 -*-
"""Generates KidWell_Pitch_Deck.docx — a styled, multi-page pitch deck."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Brand palette ────────────────────────────────────────────────
GREEN = RGBColor(0x2D, 0x9B, 0x6F)
GREEN_DARK = RGBColor(0x1F, 0x6F, 0x50)
GREEN_LIGHT = "E8F5EF"
INK = RGBColor(0x1A, 0x2E, 0x24)
MUTED = RGBColor(0x5A, 0x6F, 0x65)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
AMBER = RGBColor(0xB7, 0x79, 0x1F)
GREEN_HEX = "2D9B6F"
GREEN_DARK_HEX = "1F6F50"
PALE_HEX = "F0FAF5"

FONT = "Segoe UI"


def shade(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:color"), "auto")
    sh.set(qn("w:fill"), hex_color)
    tcPr.append(sh)


def no_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "none")
        borders.append(e)
    tblPr.append(borders)


def set_cell_margins(cell, top=100, bottom=100, left=140, right=140):
    tcPr = cell._tc.get_or_add_tcPr()
    m = OxmlElement("w:tcMar")
    for name, val in (("top", top), ("bottom", bottom), ("start", left), ("end", right)):
        node = OxmlElement(f"w:{name}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        m.append(node)
    tcPr.append(m)


def run(p, text, size=11, color=INK, bold=False, italic=False, font=FONT):
    r = p.add_run(text)
    r.font.name = font
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.bold = bold
    r.italic = italic
    return r


def space(p, before=0, after=6, line=None):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line is not None:
        pf.line_spacing = line


def add_page(doc, first=False):
    if not first:
        doc.add_page_break()


def title_bar(doc, kicker, title):
    """Colored band with a small kicker + big title."""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    no_borders(t)
    cell = t.cell(0, 0)
    shade(cell, GREEN_HEX)
    set_cell_margins(cell, top=160, bottom=160, left=220, right=220)
    cell.width = Inches(7.0)

    k = cell.paragraphs[0]
    space(k, after=2)
    run(k, kicker.upper(), size=10, color=WHITE, bold=True)

    ttl = cell.add_paragraph()
    space(ttl, after=0)
    run(ttl, title, size=22, color=WHITE, bold=True)
    doc.add_paragraph()


def h2(doc, text):
    p = doc.add_paragraph()
    space(p, before=6, after=4)
    run(p, text, size=14, color=GREEN_DARK, bold=True)
    return p


def body(doc, text, size=11, color=INK, after=6, bold=False, italic=False):
    p = doc.add_paragraph()
    space(p, after=after, line=1.15)
    run(p, text, size=size, color=color, bold=bold, italic=italic)
    return p


def bullet(doc, label, text=None):
    p = doc.add_paragraph(style="List Bullet")
    space(p, after=4, line=1.1)
    if text:
        run(p, label + ": ", size=11, color=GREEN_DARK, bold=True)
        run(p, text, size=11, color=INK)
    else:
        run(p, label, size=11, color=INK)
    return p


def stat_row(doc, stats):
    """Row of big-number stat cards. stats = [(number, label), ...]"""
    t = doc.add_table(rows=1, cols=len(stats))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    no_borders(t)
    for i, (num, label) in enumerate(stats):
        cell = t.cell(0, i)
        shade(cell, PALE_HEX)
        set_cell_margins(cell, top=180, bottom=180, left=120, right=120)
        n = cell.paragraphs[0]
        n.alignment = WD_ALIGN_PARAGRAPH.CENTER
        space(n, after=2)
        run(n, num, size=24, color=GREEN, bold=True)
        l = cell.add_paragraph()
        l.alignment = WD_ALIGN_PARAGRAPH.CENTER
        space(l, after=0)
        run(l, label, size=9.5, color=MUTED)
    doc.add_paragraph()
    return t


def info_table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    # header
    for i, htext in enumerate(headers):
        cell = t.cell(0, i)
        shade(cell, GREEN_HEX)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        space(p, after=0)
        run(p, htext, size=10.5, color=WHITE, bold=True)
    # rows
    for r_i, rowdata in enumerate(rows):
        cells = t.add_row().cells
        for i, val in enumerate(rowdata):
            cell = cells[i]
            if r_i % 2 == 1:
                shade(cell, PALE_HEX)
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            space(p, after=0, line=1.1)
            # first column bold
            run(p, val, size=10, color=INK, bold=(i == 0))
    if col_widths:
        for i, w in enumerate(col_widths):
            for cell in t.columns[i].cells:
                cell.width = Inches(w)
    return t


def callout(doc, text, fill=GREEN_LIGHT, txt_color=GREEN_DARK, bold=True, size=12):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    no_borders(t)
    cell = t.cell(0, 0)
    shade(cell, fill)
    set_cell_margins(cell, top=150, bottom=150, left=200, right=200)
    p = cell.paragraphs[0]
    space(p, after=0, line=1.15)
    run(p, text, size=size, color=txt_color, bold=bold, italic=not bold)
    doc.add_paragraph()
    return t


def footer_note(doc, n):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    space(p, before=10, after=0)
    run(p, f"KidWell  ·  Wellness Hackathon 2026  ·  {n}", size=8, color=MUTED)


# ══════════════════════════════════════════════════════════════════
doc = Document()

# Base style
normal = doc.styles["Normal"]
normal.font.name = FONT
normal.font.size = Pt(11)
normal.font.color.rgb = INK

for section in doc.sections:
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

# ── PAGE 1 — COVER ────────────────────────────────────────────────
add_page(doc, first=True)
for _ in range(3):
    doc.add_paragraph()

logo = doc.add_table(rows=1, cols=1)
logo.alignment = WD_TABLE_ALIGNMENT.CENTER
no_borders(logo)
lcell = logo.cell(0, 0)
shade(lcell, GREEN_HEX)
set_cell_margins(lcell, top=120, bottom=120, left=160, right=160)
lcell.width = Inches(1.0)
lp = lcell.paragraphs[0]
lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run(lp, "♥  KidWell", size=20, color=WHITE, bold=True)
space(lp, after=0)

doc.add_paragraph()
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
space(t, after=2)
run(t, "KidWell", size=44, color=GREEN_DARK, bold=True)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
space(sub, after=18)
run(sub, "AI-Powered School Wellbeing Platform", size=16, color=MUTED)

tag = doc.add_paragraph()
tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
space(tag, after=4)
run(tag, "Helping schools catch student wellbeing issues early —", size=13, color=INK, italic=True)
tag2 = doc.add_paragraph()
tag2.alignment = WD_ALIGN_PARAGRAPH.CENTER
space(tag2, after=30)
run(tag2, "before they become a crisis.", size=13, color=INK, italic=True)

for _ in range(4):
    doc.add_paragraph()

foot = doc.add_paragraph()
foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
space(foot, after=0)
run(foot, "Wellness Hackathon 2026  ·  Heal. Build. Thrive.", size=11, color=GREEN, bold=True)
foot2 = doc.add_paragraph()
foot2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run(foot2, "Pitch Deck", size=10, color=MUTED)

# ── PAGE 2 — THE PROBLEM ──────────────────────────────────────────
add_page(doc)
title_bar(doc, "01 · The Problem", "Schools find out too late")
body(doc, "Student wellbeing is the foundation of learning — yet most schools have no systematic way to monitor it. Stress, anxiety, poor sleep, and low mood usually go unnoticed until they erupt into a crisis.", size=12)
h2(doc, "Who feels the pain")
bullet(doc, "Students", "struggle silently; many appear fine externally while declining internally.")
bullet(doc, "Counsellors", "are responsible for hundreds of students with no early-warning signal.")
bullet(doc, "Parents", "are often the last to know something is wrong.")
h2(doc, "Why it happens")
bullet(doc, "No early signal — support is reactive, triggered only after a crisis.")
bullet(doc, "Counsellors are stretched thin and can't manually track every student.")
bullet(doc, "Existing tools are expensive, Western, and feel like surveillance.")
callout(doc, "By the time a student's struggle is visible, it is often already a crisis.")
footer_note(doc, "02")

# ── PAGE 3 — WHY IT MATTERS ───────────────────────────────────────
add_page(doc)
title_bar(doc, "02 · Why It Matters", "A large, urgent, preventable problem")
body(doc, "Mental wellbeing among young people is one of the defining challenges of our time — and the evidence shows early, preventive support works.", size=12)
stat_row(doc, [("1 in 7", "adolescents (10–19) experience a mental disorder — WHO"), ("75%", "of mental-health conditions emerge before age 24"), ("Prevention", "early intervention beats crisis response")])
h2(doc, "The opportunity in Ethiopia")
bullet(doc, "School counsellors are scarce — most students have no structured wellbeing support.")
bullet(doc, "Private urban schools have budgets, devices, and engaged parents — a ready beachhead.")
bullet(doc, "AI can extend one counsellor's reach across an entire school, affordably.")
callout(doc, "\u201cThe best solutions focus on prevention, not just crisis intervention.\u201d", fill=PALE_HEX, txt_color=GREEN_DARK, bold=False, size=12)
footer_note(doc, "03")

# ── PAGE 4 — THE SOLUTION ─────────────────────────────────────────
add_page(doc)
title_bar(doc, "03 · The Solution", "Meet KidWell")
body(doc, "KidWell is an AI-powered school wellbeing platform that turns lightweight student check-ins into early, supportive action — connecting students, parents, and counsellors on one safe interface.", size=12)
h2(doc, "One simple loop")
callout(doc, "Student checks in   →   AI spots patterns   →   Counsellor has a caring conversation", fill=GREEN_HEX, txt_color=WHITE, bold=True, size=12.5)
h2(doc, "Three connected experiences")
info_table(
    doc,
    ["Role", "What they get"],
    [
        ["Student", "A 30-second daily check-in + an instant, encouraging wellbeing score and a personalised nutrition plan."],
        ["Parent", "A clear dashboard of their child's wellbeing trend and AI summaries — positive, never alarming."],
        ["Counsellor", "A prioritised list of students who may need support, plus an AI brief for each conversation."],
    ],
    col_widths=[1.4, 5.6],
)
footer_note(doc, "04")

# ── PAGE 5 — HOW IT WORKS (AI) ────────────────────────────────────
add_page(doc)
title_bar(doc, "04 · How It Works", "Three load-bearing AI features")
body(doc, "AI is the engine, not a gimmick. Each feature does real work the product could not deliver without it.", size=12)
info_table(
    doc,
    ["AI Feature", "Input", "Output"],
    [
        ["Wellbeing Score", "Mood, energy, sleep, pain", "A 0–100 score + one encouraging, age-appropriate summary"],
        ["Nutrition Plan", "Student profile + local context", "A 3-day Ethiopian meal plan with rationale"],
        ["Counsellor Brief", "Recent check-in patterns", "Neutral summary, key concerns, and conversation starters"],
    ],
    col_widths=[1.7, 2.3, 3.0],
)
h2(doc, "Built on")
bullet(doc, "AI", "Large language model via API, called securely server-side")
bullet(doc, "App", "Next.js + React dashboard, clean and mobile-friendly")
bullet(doc, "Privacy", "Personal data is never exposed to the client or the model")
footer_note(doc, "05")

# ── PAGE 6 — SAFETY & ETHICS ──────────────────────────────────────
add_page(doc)
title_bar(doc, "05 · Safety & Ethics", "Trust is the product")
body(doc, "Children's mental-health data is the most sensitive data there is. KidWell is designed safe-by-default — and we make that our headline, not a disclaimer.", size=12)
bullet(doc, "Never diagnoses", "the AI describes patterns only — no labels, no clinical conditions.")
bullet(doc, "Human-in-the-loop", "every insight is a suggestion; a trained counsellor always decides.")
bullet(doc, "Consent-first", "parental consent and student assent before any check-in.")
bullet(doc, "Non-punitive", "wellbeing data is private to the support person, never used for discipline.")
bullet(doc, "No overreaction", "a flag appears only after multiple signals over time, never one bad day.")
callout(doc, "Not a diagnostic tool — KidWell helps humans help students, sooner.")
footer_note(doc, "06")

# ── PAGE 7 — BUSINESS MODEL ───────────────────────────────────────
add_page(doc)
title_bar(doc, "06 · Business Model", "A clear, sustainable revenue model")
body(doc, "KidWell is B2B: schools pay because it saves counsellor time and gives them a duty-of-care story for parents.", size=12)
info_table(
    doc,
    ["Stream", "Customer", "Model", "Pricing"],
    [
        ["Core SaaS", "Private schools", "Per-student / year", "~120–200 birr / student"],
        ["Premium", "Schools", "Advanced analytics add-on", "Tiered upgrade"],
        ["Expansion", "NGOs / donors", "Funded public-school rollouts", "Grant-based"],
    ],
    col_widths=[1.5, 1.9, 2.2, 1.4],
)
callout(doc, "A 500-student school ≈ 60,000–100,000 birr / year in recurring revenue.", fill=PALE_HEX, txt_color=GREEN_DARK, bold=True, size=12)
h2(doc, "Land & expand")
body(doc, "Private urban schools  →  private-school networks  →  NGO-funded public schools  →  district & government.", size=11.5, color=INK)
footer_note(doc, "07")

# ── PAGE 8 — LOCALIZATION ─────────────────────────────────────────
add_page(doc)
title_bar(doc, "07 · Localization", "Built for Ethiopia — in local languages")
body(doc, "Most wellness tools are English-only and feel foreign. KidWell is designed for Ethiopian schools, families, and counsellors from day one.", size=12)
h2(doc, "Supported languages (MVP)")
info_table(
    doc,
    ["Language", "Script", "Who it serves"],
    [
        ["English", "Latin", "International schools, demos, judges"],
        ["Amharic (አማርኛ)", "Ge'ez script", "Urban schools and families across Ethiopia"],
        ["Afaan Oromo", "Latin", "Oromia region — Ethiopia's largest language group"],
    ],
    col_widths=[2.0, 1.5, 3.5],
)
h2(doc, "What is localized")
bullet(doc, "Full UI", "login, dashboards, check-in forms, counsellor tools — switch language instantly")
bullet(doc, "AI responses", "wellbeing summaries, nutrition plans, and counsellor briefs generated in the selected language")
bullet(doc, "Cultural context", "Ethiopian foods, local diet, and school realities baked into AI prompts")
callout(doc, "One tap to switch: English · አማርኛ · Afaan Oromoo — UI and AI follow together.", fill=PALE_HEX, txt_color=GREEN_DARK, bold=True, size=12)
footer_note(doc, "08")

# ── PAGE 9 — MARKET & COMPETITION ────────────────────────────────
add_page(doc)
title_bar(doc, "08 · Market & Edge", "Why KidWell wins")
h2(doc, "Our differentiation")
bullet(doc, "Local-first", "Amharic-ready, affordable, designed for Ethiopian schools.")
bullet(doc, "Teacher-mediated", "works even where no professional counsellor exists.")
bullet(doc, "Privacy-by-design", "a trust advantage Western surveillance tools lack.")
bullet(doc, "Prevention-focused", "catches the decline early, not at crisis point.")
h2(doc, "vs. existing tools")
info_table(
    doc,
    ["", "Western platforms", "KidWell"],
    [
        ["Cost", "Expensive, enterprise", "Affordable per-student"],
        ["Context", "US/Europe-centric", "Ethiopian, local foods & language"],
        ["Stance", "Monitoring / surveillance", "Supportive, human-in-the-loop"],
        ["Counsellor", "Assumes one exists", "Works teacher-mediated too"],
    ],
    col_widths=[1.3, 2.8, 2.9],
)
footer_note(doc, "09")

# ── PAGE 10 — PRODUCT / DEMO ───────────────────────────────────────
add_page(doc)
title_bar(doc, "09 · The Product", "A working product, today")
body(doc, "KidWell is not a mockup — it is a functioning web app with three live, role-based dashboards and real AI.", size=12)
h2(doc, "What's built")
bullet(doc, "Secure role-based login", "for students, parents, and counsellors.")
bullet(doc, "Student", "daily check-in + instant AI wellbeing score; AI nutrition plan generator.")
bullet(doc, "Parent", "wellbeing overview with score trend and history.")
bullet(doc, "Counsellor", "student dashboard with risk indicators + one-click AI brief.")
h2(doc, "Demo flow (3 minutes)")
bullet(doc, "1. Student", "Sara checks in (low mood, poor sleep) → AI returns a score + encouragement.")
bullet(doc, "2. Counsellor", "opens the dashboard → Sara is flagged → generate AI brief with talking points.")
bullet(doc, "3. Close", "privacy-by-design + the business model.")
footer_note(doc, "10")

# ── PAGE 11 — ROADMAP ─────────────────────────────────────────────
add_page(doc)
title_bar(doc, "10 · Roadmap", "From hackathon to venture")
info_table(
    doc,
    ["Phase", "Focus"],
    [
        ["Now (MVP)", "Working 3-role app, AI check-in score, nutrition plan, counsellor brief"],
        ["Next 1–3 mo", "Pilot with 1 private school, parental consent flow, Amharic UI"],
        ["3–6 mo", "Trend analytics, email alerts, multi-school admin, paid subscriptions"],
        ["6–12 mo", "NGO-funded public-school rollouts, teletherapy escalation partners"],
    ],
    col_widths=[1.5, 5.5],
)
callout(doc, "This hackathon's goal: a working product + one validated school pilot.", fill=GREEN_HEX, txt_color=WHITE, bold=True, size=12)
footer_note(doc, "11")

# ── PAGE 12 — CLOSING / ASK ───────────────────────────────────────
add_page(doc)
title_bar(doc, "11 · The Ask", "Let's build healthier schools")
body(doc, "KidWell turns wellbeing from something schools react to into something they can support early, affordably, and safely.", size=12)
h2(doc, "The one-liner")
callout(doc, "We help schools support students before problems escalate — with an AI early-warning copilot that keeps humans in charge.", fill=PALE_HEX, txt_color=GREEN_DARK, bold=True, size=12.5)
h2(doc, "Why we win")
bullet(doc, "Real working AI product (not a slide).")
bullet(doc, "Clear, sustainable B2B revenue model.")
bullet(doc, "Safe-by-design — credible with educators and clinicians.")
bullet(doc, "Local-first — English, Amharic, and Afaan Oromo built in.")
doc.add_paragraph()
end = doc.add_paragraph()
end.alignment = WD_ALIGN_PARAGRAPH.CENTER
run(end, "Thank you.", size=20, color=GREEN_DARK, bold=True)
end2 = doc.add_paragraph()
end2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run(end2, "KidWell  ·  Heal. Build. Thrive.", size=11, color=GREEN, bold=True)
footer_note(doc, "11")

out = r"c:\Users\abebe\Music\Hacaton\KidWell_Pitch_Deck_v2.docx"
doc.save(out)
print("Saved:", out)
